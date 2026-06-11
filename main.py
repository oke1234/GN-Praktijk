import uuid
import json
from docx import Document
from faster_whisper import WhisperModel
from openai import OpenAI
import streamlit as st
from docx.shared import Pt

# =========================
# OPENAI
# =========================
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# =========================
# WHISPER
# =========================
@st.cache_resource
def load_whisper():
    return WhisperModel("base")

whisper_model = load_whisper()

def transcribe_audio(file_path):
    segments, _ = whisper_model.transcribe(file_path)
    return " ".join([s.text for s in segments]).strip()

# =========================
# DOCX READER
# =========================
def read_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# =========================================================
# STEP 1 — STRUCTURE EXTRACTION (STRONG JSON)
# =========================================================
def extract_blueprint(example_text):

    SYSTEM = """
You are a CLINICAL DOCUMENT STRUCTURE ENGINE.

Split the document into MULTIPLE CLEAR SECTIONS.

MANDATORY SECTIONS:
- Voeding
- Supplementen / middelen
- Afbouw / wijziging / stoppen
- Bereiding van maaltijden
- Materiaal / keukenadvies
- Overig / evaluatie (if needed)

RULES:
- NO mixing categories
- NO giant sections
- Each section must be independent
- Preserve order inside sections
- Detect bullets separately

OUTPUT ONLY JSON

FORMAT:

{
  "sections": [
    {
      "title": "",
      "blocks": [
        {
          "type": "paragraph | bullets | table",
          "content": "",
          "items": [],
          "columns": [],
          "rows": []
        }
      ]
    }
  ]
}
"""

    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[
            {"role": "system", "content": SYSTEM},
            {"role": "user", "content": example_text}
        ],
        temperature=1,
        response_format={"type": "json_object"}
    )

    return json.loads(response.choices[0].message.content)

# =========================================================
# STEP 2 — FILL SECTION (BLOCK BASED)
# =========================================================
def generate_section(section, transcript, notes, previous_consult):

    SYSTEM = """
You are a MEDICAL DOCUMENT FILL ENGINE.

RULES:
- Follow structure EXACTLY
- Fill each block correctly
- No summarization
- No merging blocks
- Keep bullets atomic (1 idea per bullet)
- Output VALID JSON for blocks only

Important
- Give each section a CLINICAL title based on content use the title 
- Put each section in their own box, do not merge content from different sections
- De info over datum, naam en volgende consult moet altijd als eerste in het document dus vooraan
"""

    USER = f"""
SECTION STRUCTURE:
{json.dumps(section, indent=2)}

TRANSCRIPT:
{transcript}

NOTES:
{notes}

PREVIOUS:
{previous_consult}

Fill this section with correct clinical content.
Return same structure, only filled content.
"""

    response = client.chat.completions.create(
        model="gpt-5-mini",
        messages=[
            {"role": "system", "content": SYSTEM},
            {"role": "user", "content": USER}
        ],
        temperature=1,
        response_format={"type": "json_object"}
    )

    return json.loads(response.choices[0].message.content)

# =========================================================
# STEP 3 — FULL DOCUMENT
# =========================================================
def generate_document(transcript, notes, example_text, previous_consult=""):

    blueprint = extract_blueprint(example_text)

    final = []

    for section in blueprint["sections"]:
        filled = generate_section(section, transcript, notes, previous_consult)
        final.append(filled)

    return final

# =========================================================
# STEP 4 — WORD EXPORT (PROFESSIONAL LAYOUT)
# =========================================================
def generate_word(sections, output_file):

    doc = Document()

    for sec in sections:

        # --- CREATE BOX (TABLE) ---
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"   # gives visible border

        cell = table.rows[0].cells[0]

        # --- SECTION TITLE INSIDE BOX ---
        p = cell.add_paragraph()
        run = p.add_run(sec["title"])
        run.bold = True
        run.font.size = Pt(16)

        # --- CONTENT INSIDE SAME BOX ---
        for block in sec["blocks"]:

            btype = block.get("type")

            if btype == "paragraph":
                cell.add_paragraph(block.get("content", ""))

            elif btype == "bullets":
                for item in block.get("items", []):
                    cell.add_paragraph(item, style="List Bullet")

            elif btype == "table":
                cols = block.get("columns", [])
                rows = block.get("rows", [])

                if cols:
                    t = cell.add_table(rows=1, cols=len(cols))

                    # header
                    for i, c in enumerate(cols):
                        t.rows[0].cells[i].text = str(c)

                    # data rows
                    for r in rows:
                        row_cells = t.add_row().cells
                        for i, val in enumerate(r):
                            row_cells[i].text = str(val)

    doc.save(output_file)
    return output_file